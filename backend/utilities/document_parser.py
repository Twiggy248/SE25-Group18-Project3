# -----------------------------------------------------------------------------
# File: document_parser.py
# Description: Document parser for ReqEngine - extracts text from various 
#              document formats including PDF, DOCX, TXT, and Markdown files.
# Author: Pradyumna Chacham
# Date: November 2025
# Copyright (c) 2025 Pradyumna Chacham. All rights reserved.
# License: MIT License - see LICENSE file in the root directory.
# -----------------------------------------------------------------------------

"""
Document Parser
Extracts text from various document formats: PDF, DOCX, TXT, MD
"""

import io
import os
import PyPDF2
from typing import Tuple
from fastapi import HTTPException, UploadFile
from docx import Document


def parse_document(content: str) -> dict:
    """
    Parse and prepare a document for requirement analysis.

    Args:
        content: Raw document content as string

    Returns:
        Dict containing parsed content and metadata
    """
    stats = get_text_stats(content)
    return {
        "text": content,
        "metadata": {
            "format": "text",
            "version": "1.0",
            "encoding": "utf-8",
            "stats": stats,
        },
    }


def extract_text_from_file(file: UploadFile) -> Tuple[str, str]:
    """
    Extract text from uploaded file

    Args:
        file: Uploaded file from FastAPI

    Returns:
        Tuple of (extracted_text, file_type)

    Raises:
        HTTPException: If file type is unsupported or parsing fails
    """
    filename = file.filename.lower()
    file_extension = os.path.splitext(filename)[1]

    # Read file content
    try:
        content = file.file.read()
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Error reading file: {str(e)}")

    # Parse based on file type
    if file_extension == ".txt" or file_extension == ".md":
        return extract_from_text(content), file_extension
    elif file_extension == ".pdf":
        return extract_from_pdf(content), file_extension
    elif file_extension == ".docx":
        return extract_from_docx(content), file_extension
    else:
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported file type: {file_extension}. Supported: .txt, .md, .pdf, .docx",
        )


def extract_from_text(content: bytes) -> str:
    """Extract text from TXT/MD files"""
    try:
        # Try UTF-8 first
        return content.decode("utf-8")
    except UnicodeDecodeError:
        # Fallback to latin-1
        try:
            return content.decode("latin-1")
        except Exception as e:
            raise HTTPException(
                status_code=400, detail=f"Error decoding text file: {str(e)}"
            )


def extract_from_pdf(content: bytes) -> str:
    """Extract text from PDF files"""

    try:
        pdf_file = io.BytesIO(content)
        pdf_reader = PyPDF2.PdfReader(pdf_file)

        text_parts = []
        for page_num, page in enumerate(pdf_reader.pages):
            try:
                text = page.extract_text()
                if text.strip():
                    text_parts.append(text)
            except Exception as e:
                print(
                    f"⚠️  Warning: Could not extract text from page {page_num + 1}: {e}"
                )
                continue

        if not text_parts:
            raise HTTPException(
                status_code=400, detail="No text could be extracted from PDF"
            )

        full_text = "\n\n".join(text_parts)
        print(
            f"✅ Extracted {len(full_text)} characters from {len(pdf_reader.pages)} pages"
        )

        return full_text

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Error parsing PDF: {str(e)}")


def extract_from_docx(content: bytes) -> str:
    """Extract text from DOCX files"""

    try:
        docx_file = io.BytesIO(content)
        doc = Document(docx_file)

        text_parts = []

        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    text_parts.append(row_text)

        if not text_parts:
            raise HTTPException(
                status_code=400, detail="No text could be extracted from DOCX"
            )

        full_text = "\n\n".join(text_parts)
        print(f"✅ Extracted {len(full_text)} characters from DOCX")

        return full_text

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Error parsing DOCX: {str(e)}")


def validate_file_size(file: UploadFile, max_size_mb: int = 10) -> None:
    """
    Validate file size

    Args:
        file: Uploaded file
        max_size_mb: Maximum allowed size in MB

    Raises:
        HTTPException: If file is too large
    """
    # Read file to check size
    content = file.file.read()
    size_mb = len(content) / (1024 * 1024)

    # Reset file pointer
    file.file.seek(0)

    if size_mb > max_size_mb:
        raise HTTPException(
            status_code=400,
            detail=f"File too large: {size_mb:.2f}MB. Maximum allowed: {max_size_mb}MB",
        )

    print(f"✅ File size: {size_mb:.2f}MB")


def get_text_stats(text: str) -> dict:
    """Get statistics about extracted text"""
    lines = text.split("\n")
    words = text.split()

    return {
        "characters": len(text),
        "words": len(words),
        "lines": len(lines),
        "estimated_tokens": len(words) * 1.3,  # Rough estimate
        "size_category": categorize_text_size(len(text)),
    }


def categorize_text_size(char_count: int) -> str:
    """Categorize text size for processing strategy"""
    if char_count < 500:
        return "tiny"  # < 500 chars
    elif char_count < 2000:
        return "small"  # 500-2K chars
    elif char_count < 8000:
        return "medium"  # 2K-8K chars
    elif char_count < 20000:
        return "large"  # 8K-20K chars
    else:
        return "very_large"  # > 20K chars
