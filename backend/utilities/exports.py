# -----------------------------------------------------------------------------
# File: export_utils.py
# Description: Export utilities for ReqEngine - exports use cases to various 
#              formats including DOCX, PlantUML, and Markdown.
# -----------------------------------------------------------------------------

"""
Export Utilities
Export use cases to various formats: DOCX, PlantUML, Markdown
"""

import os
from datetime import datetime
from typing import Dict, List, Optional
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH


def export_to_docx(
    use_cases: List[Dict], session_context: Optional[Dict], session_id: str
) -> str:
    """
    Export use cases to Microsoft Word document

    Args:
        use_cases: List of use cases
        session_context: Session metadata
        session_id: Session identifier

    Returns:
        Path to generated file
    """
    doc = Document()

    # Add title
    title = doc.add_heading("Use Case Specification", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add metadata
    if session_context:
        doc.add_paragraph(f"Project: {session_context.get('project_context', 'N/A')}")
        doc.add_paragraph(f"Domain: {session_context.get('domain', 'N/A')}")
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"Total Use Cases: {len(use_cases)}")
    doc.add_paragraph()  # Blank line

    # Add table of contents
    doc.add_heading("Table of Contents", 1)
    for idx, uc in enumerate(use_cases, 1):
        doc.add_paragraph(f"{idx}. {uc['title']}", style="List Number")

    doc.add_page_break()

    # Add each use case
    for idx, uc in enumerate(use_cases, 1):
        # Use case title
        doc.add_heading(f"{idx}. {uc['title']}", 1)

        # Preconditions
        doc.add_heading("Preconditions", 2)
        if uc.get("preconditions"):
            for precond in uc["preconditions"]:
                doc.add_paragraph(precond, style="List Bullet")
        else:
            doc.add_paragraph("None specified", style="List Bullet")

        # Main Flow
        doc.add_heading("Main Flow", 2)
        if uc.get("main_flow"):
            for i, step in enumerate(uc["main_flow"], 1):
                doc.add_paragraph(f"{i}. {step}", style="List Number")
        else:
            doc.add_paragraph("Not specified")

        # Sub Flows
        doc.add_heading("Sub Flows (Optional Paths)", 2)
        if uc.get("sub_flows") and uc["sub_flows"] != ["No subflows"]:
            for sub in uc["sub_flows"]:
                doc.add_paragraph(sub, style="List Bullet")
        else:
            doc.add_paragraph("None specified")

        # Alternate Flows
        doc.add_heading("Alternate Flows (Error Handling)", 2)
        if uc.get("alternate_flows") and uc["alternate_flows"] != [
            "No alternate flows"
        ]:
            for alt in uc["alternate_flows"]:
                doc.add_paragraph(alt, style="List Bullet")
        else:
            doc.add_paragraph("None specified")

        # Outcomes
        doc.add_heading("Expected Outcomes", 2)
        if uc.get("outcomes"):
            for outcome in uc["outcomes"]:
                doc.add_paragraph(outcome, style="List Bullet")
        else:
            doc.add_paragraph("Not specified")

        # Stakeholders
        doc.add_heading("Stakeholders", 2)
        if uc.get("stakeholders"):
            stakeholder_text = ", ".join(uc["stakeholders"])
            doc.add_paragraph(stakeholder_text)
        else:
            doc.add_paragraph("Not specified")

        # Add page break between use cases (except last one)
        if idx < len(use_cases):
            doc.add_page_break()

    # Save document
    export_dir = "/tmp"
    os.makedirs(export_dir, exist_ok=True)
    file_path = os.path.join(export_dir, f"use_cases_{session_id}.docx")
    doc.save(file_path)

    return file_path


def export_to_plantuml(use_cases: List[Dict]) -> str:
    """
    Export use cases as PlantUML diagram

    Args:
        use_cases: List of use cases

    Returns:
        PlantUML code as string
    """
    plantuml = "@startuml\n"
    plantuml += "left to right direction\n"
    plantuml += "skinparam packageStyle rectangle\n\n"

    def clean_id(text):
        """Create valid PlantUML identifier by removing/replacing special characters"""
        return (
            text.replace(" ", "_")
            .replace("-", "_")
            .replace("/", "_")
            .replace("(", "")
            .replace(")", "")
            .replace("!", "")
        )

    # Extract all unique stakeholders (actors)
    actors = set()
    for uc in use_cases:
        if uc.get("stakeholders"):
            actors.update(uc["stakeholders"])

    # Add actors
    plantuml += "' Actors\n"
    actor_map = {}
    for actor in sorted(actors):
        # Create valid PlantUML identifier and display name
        actor_id = clean_id(actor)
        actor_map[actor] = actor_id

        # Use different notation for system vs human actors
        if (
            "system" in actor.lower()
            or "database" in actor.lower()
            or "api" in actor.lower()
            or "db" in actor.lower()
        ):
            plantuml += f'rectangle {actor_id} as "{actor_id}"\n'
        else:
            plantuml += f'actor {actor_id} as "{actor_id}"\n'

    plantuml += "\n"

    # Add use cases
    plantuml += "' Use Cases\n"
    uc_map = {}
    for idx, uc in enumerate(use_cases):
        uc_id = clean_id(f"UC{idx + 1}")
        uc_map[uc["title"]] = uc_id

        # Escape quotes and clean title for PlantUML
        title = uc["title"].replace('"', '\\"')
        clean_title = clean_id(title)
        plantuml += f'usecase {uc_id} as "{clean_title}"\n'

    plantuml += "\n"

    # Connect actors to use cases
    plantuml += "' Relationships\n"
    for uc in use_cases:
        uc_id = uc_map[uc["title"]]
        if uc.get("stakeholders"):
            for stakeholder in uc["stakeholders"]:
                actor_id = actor_map.get(stakeholder)
                if actor_id:
                    plantuml += f"{actor_id} --> {uc_id}\n"

    # Add dependencies between use cases if they share flows
    plantuml += "\n' Use Case Dependencies (include/extend)\n"
    for i, uc1 in enumerate(use_cases):
        uc1_id = uc_map[uc1["title"]]
        for j, uc2 in enumerate(use_cases):
            if i >= j:
                continue

            uc2_id = uc_map[uc2["title"]]

            # Check if uc2 is mentioned in uc1's flows
            uc1_text = " ".join(uc1.get("main_flow", []) + uc1.get("sub_flows", []))
            if uc2["title"].lower() in uc1_text.lower():
                plantuml += f"{uc1_id} ..> {uc2_id} : <<include>>\n"

    plantuml += "\n@enduml"

    return plantuml


def export_to_markdown(
    use_cases: List[Dict], session_context: Optional[Dict], session_id: str
) -> str:
    """
    Export use cases to Markdown document

    Args:
        use_cases: List of use cases
        session_context: Session metadata
        session_id: Session identifier

    Returns:
        Path to generated file
    """
    md = "# Use Case Specification\n\n"

    # Add metadata
    if session_context:
        md += f"**Project:** {session_context.get('project_context', 'N/A')}  \n"
        md += f"**Domain:** {session_context.get('domain', 'N/A')}  \n"
    md += f"**Generated:** {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}  \n"
    md += f"**Total Use Cases:** {len(use_cases)}  \n\n"

    # Add table of contents
    md += "## Table of Contents\n\n"
    for idx, uc in enumerate(use_cases, 1):
        # Create anchor link
        anchor = uc["title"].lower().replace(" ", "-").replace("/", "-")
        md += f"{idx}. [{uc['title']}](#{anchor})\n"
    md += "\n---\n\n"

    # Add each use case
    for idx, uc in enumerate(use_cases, 1):
        md += f"## {idx}. {uc['title']}\n\n"

        # Preconditions
        md += "### Preconditions\n\n"
        if uc.get("preconditions") and uc["preconditions"] != ["No preconditions"]:
            for precond in uc["preconditions"]:
                md += f"- {precond}\n"
        else:
            md += "*None specified*\n"
        md += "\n"

        # Main Flow
        md += "### Main Flow\n\n"
        if uc.get("main_flow") and uc["main_flow"] != ["No main flow"]:
            for i, step in enumerate(uc["main_flow"], 1):
                md += f"{i}. {step}\n"
        else:
            md += "*Not specified*\n"
        md += "\n"

        # Sub Flows
        md += "### Sub Flows (Optional Paths)\n\n"
        if uc.get("sub_flows") and uc["sub_flows"] != ["No subflows"]:
            for sub in uc["sub_flows"]:
                md += f"- {sub}\n"
        else:
            md += "*None specified*\n"
        md += "\n"

        # Alternate Flows
        md += "### Alternate Flows (Error Handling)\n\n"
        if uc.get("alternate_flows") and uc["alternate_flows"] != [
            "No alternate flows"
        ]:
            for alt in uc["alternate_flows"]:
                md += f"- {alt}\n"
        else:
            md += "*None specified*\n"
        md += "\n"

        # Outcomes
        md += "### Expected Outcomes\n\n"
        if uc.get("outcomes") and uc["outcomes"] != ["No outcomes"]:
            for outcome in uc["outcomes"]:
                md += f"- {outcome}\n"
        else:
            md += "*Not specified*\n"
        md += "\n"

        # Stakeholders
        md += "### Stakeholders\n\n"
        if uc.get("stakeholders") and uc["stakeholders"] != ["No stakeholders"]:
            stakeholder_list = ", ".join(uc["stakeholders"])
            md += f"{stakeholder_list}\n"
        else:
            md += "*Not specified*\n"
        md += "\n"

        # Separator between use cases
        if idx < len(use_cases):
            md += "---\n\n"

    # Save markdown file
    export_dir = "/tmp"
    os.makedirs(export_dir, exist_ok=True)
    file_path = os.path.join(export_dir, f"use_cases_{session_id}.md")

    with open(file_path, "w", encoding="utf-8") as f:
        f.write(md)

    return file_path


def export_to_html(use_cases: List[Dict]) -> str:
    """
    Convert use cases to HTML format

    Args:
        use_cases: List of use cases to convert

    Returns:
        HTML formatted string containing use case documentation
    """
    if not use_cases:
        return (
            "<!DOCTYPE html><html><body><h1>No Use Cases Available</h1></body></html>"
        )

    html = [
        "<!DOCTYPE html><html><head><style>",
        "body { font-family: Arial, sans-serif; margin: 2em; }",
        "h1 { color: #333; }",
        "h2 { color: #666; }",
        "ul { margin-left: 2em; }",
        ".use-case { border: 1px solid #ddd; padding: 1em; margin: 1em 0; border-radius: 4px; }",
        "</style></head><body>",
        "<h1>Use Cases Documentation</h1>",
    ]

    for uc in use_cases:
        html.extend(
            [
                f'<div class="use-case">',
                f'<h2>{uc.get("title", "Untitled")}</h2>',
                "<h3>Actor</h3>",
                f'<p>{uc.get("actor", "Not specified")}</p>',
                "<h3>Goal</h3>",
                f'<p>{uc.get("goal", "Not specified")}</p>',
                "<h3>Preconditions</h3>",
                "<ul>",
            ]
        )

        for precond in uc.get("preconditions", []):
            html.append(f"<li>{precond}</li>")

        html.extend(["</ul>", "<h3>Main Flow</h3>", "<ol>"])
        for step in uc.get("main_flow", []):
            html.append(f"<li>{step}</li>")

        html.extend(["</ol>", "<h3>Alternative Flows</h3>", "<ul>"])
        for flow in uc.get("alternate_flows", []):
            html.append(f"<li>{flow}</li>")

        html.extend(["</ul>", "<h3>Outcomes</h3>", "<ul>"])
        for outcome in uc.get("outcomes", []):
            html.append(f"<li>{outcome}</li>")

        html.append("</ul></div>")

    html.append("</body></html>")
    return "\n".join(html)


def export_to_format(use_cases: List[Dict], format_type: str) -> Dict:
    """
    Export use cases to specified format.

    Args:
        use_cases: List of use cases
        format_type: Target format (json, docx, plantuml, markdown, jira, html)

    Returns:
        Dict with export results and metadata
    """
    if format_type.lower() == "json":
        result = export_to_json(use_cases, None)
        return {"status": "success", "formats": ["JSON"], "data": result}
    elif format_type.lower() == "docx":
        path = export_to_docx(use_cases, None, "default")
        return {"status": "success", "formats": ["DOCX"], "export_path": str(path)}
    elif format_type.lower() == "plantuml":
        result = export_to_plantuml(use_cases)
        return {"status": "success", "formats": ["PlantUML"], "data": result}
    elif format_type.lower() == "markdown":
        path = export_to_markdown(use_cases, None, "default")
        return {"status": "success", "formats": ["Markdown"], "export_path": str(path)}
    elif format_type.lower() == "jira":
        return {
            "status": "success",
            "formats": ["JIRA"],
            "data": {"issues": [_convert_to_jira_issue(uc) for uc in use_cases]},
        }
    elif format_type.lower() == "html":
        return {
            "status": "success",
            "formats": ["HTML"],
            "data": export_to_html(use_cases),
        }
    else:
        raise ValueError(f"Unsupported format: {format_type}")

def export_to_json(use_cases: List[Dict], session_context: Optional[Dict]) -> Dict:
    """
    Export use cases to structured JSON format

    Args:
        use_cases: List of use cases
        session_context: Session metadata

    Returns:
        Structured JSON object
    """
    return {
        "metadata": {
            "project_context": (
                session_context.get("project_context", "") if session_context else ""
            ),
            "domain": session_context.get("domain", "") if session_context else "",
            "generated_at": datetime.now().isoformat(),
            "total_use_cases": len(use_cases),
        },
        "use_cases": use_cases,
    }

def _convert_to_jira_issue(use_case: Dict) -> Dict:
    """Convert use case to JIRA issue format"""
    return {
        "summary": use_case.get("title", "Untitled Use Case"),
        "description": _build_jira_description(use_case),
        "issuetype": "Story",
        "labels": ["use-case"],
    }


def _build_jira_description(use_case: Dict) -> str:
    """Build JIRA description from use case"""
    parts = []

    if "preconditions" in use_case:
        parts.append("h3. Preconditions")
        for pre in use_case["preconditions"]:
            parts.append(f"* {pre}")

    if "main_flow" in use_case:
        parts.append("\nh3. Main Flow")
        for step in use_case["main_flow"]:
            parts.append(f"# {step}")

    if "requirements" in use_case:
        parts.append("\nh3. Requirements")
        for req in use_case["requirements"]:
            parts.append(f"* {req}")

    return "\n".join(parts)
